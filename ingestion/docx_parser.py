"""
DOCX Parser — extracts structured blocks from Word documents.

Uses python-docx to preserve:
- Heading styles (Heading 1, Heading 2, etc.)
- List items (detected by paragraph style or numbering)
- Tables
- Normal paragraphs
- Code blocks (detected by monospace style or code fence markers)

Also strips Word tracked-changes markup before extraction:
- w:del elements (deleted text) are removed entirely
- w:ins elements (inserted text) are unwrapped — content kept, wrapper removed
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from models.canonical import DocumentBlock, ParsedDocument

# Reuse the text helpers from pdf_parser
from ingestion.pdf_parser import _blocks_to_text

# Word namespace
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def parse(file_path: str) -> ParsedDocument:
    """
    Parse a .docx file into a ParsedDocument.

    Tracked changes are stripped before content extraction:
    deleted text (w:del) is removed, inserted text (w:ins) is kept.

    Args:
        file_path: Path to the .docx file.

    Returns:
        ParsedDocument with typed DocumentBlocks.
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install with: pip install python-docx"
        )

    path = Path(file_path)
    document = docx.Document(file_path)

    # Strip tracked changes from the XML before processing
    _strip_tracked_changes(document.element.body)

    blocks: List[DocumentBlock] = []
    title = _extract_title(document) or _filename_to_title(path.stem)

    # --- Process document body ---
    for element in document.element.body:
        tag = element.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            para = _find_paragraph(document, element)
            if para is None:
                continue
            block = _process_paragraph(para)
            if block:
                blocks.append(block)

        elif tag == "tbl":
            table = _find_table(document, element)
            if table is None:
                continue
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                cell_text = " | ".join(c for c in cells if c)
                if cell_text.strip():
                    blocks.append(DocumentBlock(
                        block_type="table_row",
                        content=cell_text,
                        row_index=row_idx,
                    ))

    full_text = _blocks_to_text(blocks)
    return ParsedDocument(
        title=title,
        source_file=file_path,
        source_format="docx",
        blocks=blocks,
        full_text=full_text,
    )


# ---------------------------------------------------------------------------
# Tracked-changes stripping
# ---------------------------------------------------------------------------

def _strip_tracked_changes(body_element) -> None:
    """
    Remove Word tracked-changes markup in-place from the document body element.

    - w:del  → entire element removed (deleted text not shown)
    - w:ins  → wrapper element removed, children promoted to parent (text kept)

    This operates directly on the lxml element tree so subsequent docx parsing
    sees a clean document with only the accepted/final content.
    """
    del_tag = f"{{{_W_NS}}}del"
    ins_tag = f"{{{_W_NS}}}ins"

    # Iterate over all descendants; collect actions first to avoid modifying
    # the tree while iterating.
    to_remove: list = []
    to_unwrap: list = []

    for elem in body_element.iter():
        if elem.tag == del_tag:
            to_remove.append(elem)
        elif elem.tag == ins_tag:
            to_unwrap.append(elem)

    # Remove deleted nodes (skip if already removed by a parent removal)
    removed: set = set()
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None and id(elem) not in removed:
            idx = list(parent).index(elem)
            parent.remove(elem)
            removed.add(id(elem))

    # Unwrap inserted nodes: promote children, then remove the w:ins wrapper
    for elem in to_unwrap:
        parent = elem.getparent()
        if parent is None or id(elem) in removed:
            continue
        idx = list(parent).index(elem)
        # Move children before the w:ins element
        for child in list(elem):
            parent.insert(idx, child)
            idx += 1
        parent.remove(elem)
        removed.add(id(elem))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _process_paragraph(para) -> DocumentBlock | None:
    text = para.text.strip()
    if not text:
        return None

    style_name = para.style.name if para.style else ""

    # Heading detection
    heading_match = re.match(r"Heading\s*(\d+)", style_name, re.IGNORECASE)
    if heading_match:
        level = int(heading_match.group(1))
        return DocumentBlock(block_type="heading", content=text, level=level)

    # Title style
    if "title" in style_name.lower():
        return DocumentBlock(block_type="heading", content=text, level=0)

    # List Item styles
    if "list" in style_name.lower() or _is_list_paragraph(para):
        level = _get_list_level(para)
        clean = _strip_list_prefix(text)
        return DocumentBlock(block_type="list_item", content=clean, level=level + 1)

    # Code / verbatim style
    if any(kw in style_name.lower() for kw in ("code", "verbatim", "mono", "preformat")):
        return DocumentBlock(block_type="code_block", content=text, level=0)

    # Numbered list heuristic (for docs without proper styles)
    if re.match(r"^\s*\d+[\.\)]\s+", text) or re.match(r"^\s*Step\s+\d+", text, re.IGNORECASE):
        clean = re.sub(r"^\s*(\d+[\.\)]\s+|Step\s+\d+\s*[-:.]?\s*)", "", text, flags=re.IGNORECASE).strip()
        return DocumentBlock(block_type="list_item", content=clean, level=1)

    # Bullet heuristic
    if re.match(r"^\s*[-•◦▪▸*]\s+", text):
        clean = re.sub(r"^\s*[-•◦▪▸*]\s+", "", text).strip()
        return DocumentBlock(block_type="list_item", content=clean, level=1)

    return DocumentBlock(block_type="paragraph", content=text, level=0)


def _is_list_paragraph(para) -> bool:
    """Check if paragraph has numbering XML element (Word list numbering)."""
    try:
        return para._p.pPr is not None and para._p.pPr.numPr is not None
    except Exception:
        return False


def _get_list_level(para) -> int:
    """Get the indentation level of a list paragraph."""
    try:
        num_pr = para._p.pPr.numPr
        if num_pr is not None and num_pr.ilvl is not None:
            return int(num_pr.ilvl.val)
    except Exception:
        pass
    return 0


def _strip_list_prefix(text: str) -> str:
    text = re.sub(r"^\s*\d+[\.\)]\s+", "", text)
    text = re.sub(r"^\s*[-•◦▪▸*]\s+", "", text)
    return text.strip()


def _extract_title(document) -> str | None:
    """Try to get the document title from core properties or first heading."""
    try:
        core = document.core_properties
        if core.title:
            return core.title.strip()
    except Exception:
        pass

    # Fall back to first heading paragraph
    for para in document.paragraphs:
        if para.style and "heading" in para.style.name.lower():
            text = para.text.strip()
            if text:
                return text

    return None


def _filename_to_title(stem: str) -> str:
    return stem.replace("_", " ").replace("-", " ").title()


def _find_paragraph(document, element):
    """Find the docx Paragraph object for an lxml element."""
    import docx.oxml.ns as ns
    for para in document.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(document, element):
    """Find the docx Table object for an lxml element."""
    for table in document.tables:
        if table._element is element:
            return table
    return None
